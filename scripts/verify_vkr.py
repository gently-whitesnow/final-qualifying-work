"""Сверка собранной ВКР: реферат и содержание ↔ фактический PDF.

Запускать после `build_vkr_docx.py` и рендера. Скрипт ничего не правит,
только сообщает о расхождениях. Выход 0 — всё совпадает, 1 — есть расхождения.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parent.parent
DOCX = ROOT / "build/docx/vkr-draft-1.docx"
PDF = ROOT / "build/docx/rendered/vkr-draft-1.pdf"

# Печатные страницы нумеруются с третьей физической: титульный лист
# и страница задания идут без видимых номеров.
PRINT_PAGE_OFFSET = 2

TOC_ITEMS = [
    "ЗАКЛЮЧЕНИЕ",
    "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
    "ПРИЛОЖЕНИЕ А",
    "ПРИЛОЖЕНИЕ Б",
    "ПРИЛОЖЕНИЕ В",
]


def extract_abstract_counts(doc: Document) -> dict[str, int]:
    text = ""
    for p in doc.paragraphs:
        if "содержит" in p.text and "страниц" in p.text:
            text = p.text
            break
    if not text:
        raise SystemExit("Не нашёл фразу 'содержит N страниц' в реферате")

    def grab(pattern: str) -> int:
        m = re.search(pattern, text)
        if not m:
            raise SystemExit(f"Не нашёл паттерн {pattern!r} в реферате")
        return int(m.group(1))

    return {
        "pages": grab(r"содержит\s+(\d+)\s+страниц"),
        "figures": grab(r"(\d+)\s+рисун"),
        "tables": grab(r"(\d+)\s+таблиц"),
        "sources": grab(r"из\s+(\d+)\s+наименований"),
        "appendices": grab(r"и\s+(\d+)\s+приложени"),
    }


def count_in_docx(doc: Document) -> dict[str, int]:
    figs = sum(1 for p in doc.paragraphs if p.text.startswith("Рисунок ") and "—" in p.text)
    tabs = sum(1 for p in doc.paragraphs if p.text.startswith("Таблица ") and "—" in p.text)

    hits = [i for i, p in enumerate(doc.paragraphs) if "СПИСОК ИСПОЛЬЗОВАННЫХ" in p.text]
    src_count = 0
    if len(hits) >= 2:
        start = hits[-1]
        end = next(
            (i for i, p in enumerate(doc.paragraphs) if i > start and p.text.strip().startswith("ПРИЛОЖЕНИЕ")),
            len(doc.paragraphs),
        )
        src_count = sum(
            1 for i in range(start + 1, end) if re.match(r"^\d+\.\s", doc.paragraphs[i].text.strip())
        )

    appendix_letters: set[str] = set()
    for p in doc.paragraphs:
        m = re.match(r"^ПРИЛОЖЕНИЕ\s+(\S+)", p.text.strip())
        if m:
            appendix_letters.add(m.group(1))
    appendix_count = len(appendix_letters)

    return {"figures": figs, "tables": tabs, "sources": src_count, "appendices": appendix_count}


def extract_toc_pages(doc: Document) -> dict[str, int]:
    toc: dict[str, int] = {}
    for p in doc.paragraphs:
        t = p.text.strip()
        for item in TOC_ITEMS:
            if t.startswith(item) and item not in toc:
                m = re.search(r"(\d+)$", t)
                if m:
                    toc[item] = int(m.group(1))
    return toc


def find_actual_pages(pdf: PdfReader) -> dict[str, int]:
    found: dict[str, int] = {}
    for i, page in enumerate(pdf.pages, 1):
        text = page.extract_text() or ""
        # пропускаем страницу содержания, где встречаются все заголовки разом
        if "СОДЕРЖАНИЕ" in text[:80]:
            continue
        for item in TOC_ITEMS:
            if item not in found and item in text:
                found[item] = i - PRINT_PAGE_OFFSET
    return found


def main() -> int:
    if not DOCX.exists():
        print(f"DOCX не найден: {DOCX}", file=sys.stderr)
        return 2
    if not PDF.exists():
        print(f"PDF не найден: {PDF}", file=sys.stderr)
        return 2

    doc = Document(str(DOCX))
    pdf = PdfReader(str(PDF))

    promised = extract_abstract_counts(doc)
    actual_counts = count_in_docx(doc)
    actual_counts["pages"] = len(pdf.pages)
    toc_promised = extract_toc_pages(doc)
    toc_actual = find_actual_pages(pdf)

    failures: list[str] = []

    for key in ("pages", "figures", "tables", "sources", "appendices"):
        if promised[key] != actual_counts[key]:
            failures.append(
                f"реферат говорит {promised[key]} {key}, фактически {actual_counts[key]}"
            )

    for item in TOC_ITEMS:
        if item not in toc_promised:
            failures.append(f"в содержании нет пункта {item!r}")
            continue
        if item not in toc_actual:
            failures.append(f"в PDF не найдено вхождение {item!r}")
            continue
        if toc_promised[item] != toc_actual[item]:
            failures.append(
                f"содержание: {item} = стр. {toc_promised[item]}, "
                f"фактически = стр. {toc_actual[item]}"
            )

    if failures:
        print("Расхождения:")
        for f in failures:
            print(f"  - {f}")
        return 1

    print("ВКР целостна:")
    print(f"  страницы: {actual_counts['pages']}")
    print(f"  рисунки: {actual_counts['figures']}, таблицы: {actual_counts['tables']}")
    print(f"  источники: {actual_counts['sources']}, приложения: {actual_counts['appendices']}")
    print(f"  страницы пунктов содержания: {toc_actual}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
