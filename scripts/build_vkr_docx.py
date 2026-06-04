from __future__ import annotations

import re
import sys
import textwrap
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "materials" / "04-docx" / "vkr-draft-content.md"
OUT_DIR = ROOT / "build" / "docx"
OUT_DOCX = OUT_DIR / "vkr-draft-1.docx"
PRACTICE_OUT_DOCX = OUT_DIR / "practice-report-draft-1.docx"
ASSET_DIR = OUT_DIR / "assets"

TITLE = "Разработка real-time интерфейса с поддержкой горизонтального масштабирования WebSocket-соединений для системы отслеживания задач"
PRACTICE_TYPE = "производственная"
PRACTICE_KIND = "преддипломная"
STUDENT_FULL = "Зайцев Александр Сергеевич"
STUDENT_INSTR = "Зайцевым Александром Сергеевичем"
STUDENT_SHORT = "А.С. Зайцев"
GROUP = "4131з"
VKR_SUPERVISOR = "А.В. Фомин"
VKR_SUPERVISOR_FULL = "Фомин Александр Владимирович"
VKR_SUPERVISOR_ROLE = "канд. техн. наук, доцент"
PRACTICE_SUPERVISOR = "С.А. Рогачев"
PRACTICE_SUPERVISOR_ROLE = "ст. преподаватель"
CITY_YEAR = "Санкт-Петербург 2026"
FONT_REGULAR = Path("/System/Library/Fonts/Supplemental/Times New Roman.ttf")
FONT_BOLD = Path("/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf")

VKR_TITLE_TEMPLATE = ROOT / "materials" / "00-source-documents" / "Бакалавриат(ЗАОЧНОЕ)_Бланк_ТитЛИСТ_ЗаданиеВКРБ_2025.docx"
VKR_GOAL = (
    "разработка real-time интерфейса для системы отслеживания задач с поддержкой "
    "горизонтального масштабирования WebSocket-соединений."
)
VKR_TASKS = (
    "1) проанализировать предметную область; 2) спроектировать распределенную архитектуру "
    "real-time контура; 3) реализовать решение; 4) провести испытания."
)
VKR_CONTENT_SECTIONS = (
    "введение, анализ предметной области, проектирование real-time контура, "
    "разработка и тестирование, заключение."
)
VKR_GOAL_LINES = [
    "разработка real-time интерфейса для системы отслеживания задач с поддержкой",
    "горизонтального масштабирования WebSocket-соединений.",
    "",
]
VKR_TASK_LINES = [
    "1) проанализировать предметную область; 2) спроектировать",
    "распределенную архитектуру real-time контура; 3) реализовать решение; 4) провести испытания.",
    "",
]
VKR_CONTENT_LINES = [
    "введение, анализ предметной области,",
    "проектирование real-time контура, разработка и тестирование, заключение.",
    "",
    "",
]


_INLINE_CODE_RE = re.compile(r"(`[^`]+`)")

PRACTICE_LITERATURE_REVIEW = """
### 1.1 Описание использованной литературы

При подготовке отчета использовались нормативные документы, официальная техническая документация и учебные издания по программной инженерии, сетевым технологиям, базам данных и тестированию. Источники подбирались так, чтобы обосновать как предметную часть работы, так и выбранный способ реализации масштабируемого real-time контура.

1. RFC 6455 описывает протокол WebSocket и используется как базовый источник по сетевому механизму двустороннего обмена данными между клиентом и сервером.
2. Документация Microsoft по ASP.NET Core SignalR применяется для описания хабов, клиентских событий, групп и общей модели real-time взаимодействия.
3. Документация Microsoft по размещению и масштабированию SignalR используется при рассмотрении ограничений single-node архитектуры и условий работы в многосерверной конфигурации.
4. Документация Microsoft по Redis backplane для SignalR является основным источником при обосновании выбранного подхода к межузловой доставке событий.
5. Документация Microsoft по пользователям и группам SignalR применяется для описания модели подписки клиента на группу отчета и необходимости повторного вступления в группу после переподключения.
6. Документация Microsoft по JavaScript-клиенту SignalR используется при описании проверочного клиента, WebSocket-подключения и автоматического восстановления соединения.
7. Документация Microsoft по конфигурации SignalR применяется при описании настроек транспорта, подключения клиента и параметров, влияющих на проверочный стенд.
8. Документация Microsoft по диагностике SignalR используется для обоснования добавленного диагностического контура, логирования и фиксации идентификаторов соединений.
9. Документация Redis Pub/Sub используется для понимания роли Redis как механизма публикации и подписки при межузловой доставке transient UI-событий.
10. Документация nginx по WebSocket proxying применяется при описании reverse proxy, передачи заголовков Upgrade и балансировки WebSocket-трафика.
11. ГОСТ 7.32-2017 используется как нормативная основа структуры и оформления отчета.
12. ГОСТ Р 7.0.100-2018 используется как нормативная основа оформления библиографического списка.
13. Учебно-методическое пособие С.А. Рогачева и Ю. Бабюк используется для учета требований кафедры к отчету по преддипломной практике.
14. Учебник В.Г. Олифера и Н.А. Олифера применяется для общей сетевой терминологии и понимания принципов взаимодействия сетевых приложений.
15. Учебник С.А. Орлова используется для связи практической разработки с понятиями программной инженерии и жизненного цикла программного продукта.
16. Учебное пособие Л.Г. Гагариной, Е.В. Кокоревой и В.Д. Виснадул применяется при описании разработки программного обеспечения и фиксации требований.
17. Учебник Б.Я. Советова, В.В. Цехановского и В.Д. Чертовского используется как источник по роли базы данных в составе информационной системы.
18. Учебник В.А. Гвоздевой и Б.А. Баллода применяется при описании систем отслеживания задач как разновидности автоматизированной информационной системы.
19. Учебник В.В. Липаева используется для обоснования инженерного подхода к разработке, сопровождению и качеству программных систем.
20. Книга С.С. Куликова используется при описании проверки результата, тестовых сценариев и интерпретации результатов испытаний.

Таким образом, использованная литература покрывает четыре смысловые группы: нормативные требования к отчету, теорию WebSocket и SignalR, инфраструктурные средства масштабирования и общие положения программной инженерии. Это позволяет рассматривать выполненную работу не только как настройку стенда, но и как инженерную разработку, связанную с архитектурой, реализацией и проверкой программной системы.
"""


def add_inline_runs(paragraph, text: str, *, size: int = 14, bold: bool = False, italic: bool = False):
    if not text:
        return
    for part in _INLINE_CODE_RE.split(text):
        if not part:
            continue
        is_code = part.startswith("`") and part.endswith("`") and len(part) >= 2
        run = paragraph.add_run(part[1:-1] if is_code else part)
        set_run_font(
            run,
            size=size,
            bold=bold,
            italic=italic,
            name="Courier New" if is_code else "Times New Roman",
        )


def set_cell_text(cell, text: str, *, size: int = 14, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    add_inline_runs(p, text, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def pil_font(size: int, *, bold: bool = False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(str(path), size=size)


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        width = draw.textbbox((0, 0), candidate, font=font)[2]
        if width <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int, int, int],
    text: str,
    *,
    fill: str = "#F7F7F7",
    outline: str = "#222222",
    font=None,
    bold: bool = False,
):
    x1, y1, x2, y2 = xy
    font = font or pil_font(25, bold=bold)
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    lines = wrap_text(draw, text, font, x2 - x1 - 28)
    line_height = font.size + 6
    total_height = line_height * len(lines)
    y = y1 + ((y2 - y1) - total_height) // 2
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, fill="#111111", font=font)
        y += line_height


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, color: str = "#111111", width: int = 4):
    draw.line([start, end], fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        head = [(x2, y2), (x2 - sign * 18, y2 - 10), (x2 - sign * 18, y2 + 10)]
    else:
        sign = 1 if y2 >= y1 else -1
        head = [(x2, y2), (x2 - 10, y2 - sign * 18), (x2 + 10, y2 - sign * 18)]
    draw.polygon(head, fill=color)


def add_figure_title(draw: ImageDraw.ImageDraw, title: str):
    font = pil_font(32, bold=True)
    bbox = draw.textbbox((0, 0), title, font=font)
    draw.text(((1600 - (bbox[2] - bbox[0])) // 2, 30), title, fill="#111111", font=font)


def draw_author_marker(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str = "Разработано в рамках ВКР"):
    x1, y1, x2, y2 = xy
    font = pil_font(19, bold=True)
    draw.rounded_rectangle(xy, radius=12, fill="#E8F3F1", outline="#0A6E73", width=3)
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y1 + ((y2 - y1) - (bbox[3] - bbox[1])) // 2 - 2), text, fill="#0A4C4C", font=font)


def create_single_node_figure(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    add_figure_title(draw, "Исходная одноузловая архитектура")
    node_fill = "#EEF3FA"
    service_fill = "#F9F0DF"
    data_fill = "#EAF6EA"
    draw_box(draw, (90, 170, 330, 280), "Клиент A", fill=node_fill)
    draw_box(draw, (90, 420, 330, 530), "Клиент B", fill=node_fill)
    draw_box(draw, (470, 300, 710, 410), "nginx", fill="#F3F3F3", bold=True)
    draw_box(draw, (850, 300, 1130, 410), "app-api", fill=service_fill, bold=True)
    draw_box(draw, (1210, 270, 1510, 420), "Локальные соединения и группы SignalR", fill="#FCECEB")
    draw_box(draw, (1210, 570, 1510, 690), "PostgreSQL", fill=data_fill)
    draw_arrow(draw, (330, 225), (470, 335))
    draw_arrow(draw, (330, 475), (470, 375))
    draw_arrow(draw, (710, 355), (850, 355))
    draw_arrow(draw, (1130, 350), (1210, 345))
    draw_arrow(draw, (1130, 390), (1210, 630))
    note_font = pil_font(24)
    note = "Состояние соединений и групп хранится в памяти одного серверного процесса."
    draw.rounded_rectangle((120, 760, 1480, 835), radius=14, fill="#FFF8D6", outline="#D2B43A", width=2)
    draw.text((170, 785), note, fill="#111111", font=note_font)
    img.save(path)


def create_no_backplane_figure(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    add_figure_title(draw, "Несколько экземпляров без Redis backplane")
    node_fill = "#EEF3FA"
    service_fill = "#F9F0DF"
    group_fill = "#FCECEB"
    draw_box(draw, (80, 170, 320, 280), "Клиент A", fill=node_fill)
    draw_box(draw, (80, 500, 320, 610), "Клиент B", fill=node_fill)
    draw_box(draw, (440, 330, 680, 450), "nginx", fill="#F3F3F3", bold=True)
    draw_box(draw, (810, 160, 1100, 280), "app-api-1", fill=service_fill, bold=True)
    draw_box(draw, (810, 520, 1100, 640), "app-api-2", fill=service_fill, bold=True)
    draw_box(draw, (1220, 130, 1510, 310), "Локальные группы SignalR узла 1", fill=group_fill)
    draw_box(draw, (1220, 490, 1510, 670), "Локальные группы SignalR узла 2", fill=group_fill)
    draw_arrow(draw, (320, 225), (440, 360))
    draw_arrow(draw, (320, 555), (440, 420))
    draw_arrow(draw, (680, 360), (810, 220))
    draw_arrow(draw, (680, 420), (810, 580))
    draw_arrow(draw, (1100, 220), (1220, 220), color="#0A6E73")
    draw_arrow(draw, (1100, 580), (1220, 580), color="#0A6E73")
    draw.line((1110, 330, 1210, 470), fill="#C0392B", width=8)
    draw.line((1110, 470, 1210, 330), fill="#C0392B", width=8)
    small = pil_font(24, bold=True)
    note = "Событие, созданное на app-api-1, не попадает в локальные группы app-api-2."
    draw.rounded_rectangle((120, 770, 1480, 850), radius=14, fill="#FFF3D8", outline="#C76F2E", width=3)
    bbox = draw.textbbox((0, 0), note, font=small)
    draw.text(((1600 - (bbox[2] - bbox[0])) // 2, 797), note, fill="#111111", font=small)
    img.save(path)


def create_scaleout_figure(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    add_figure_title(draw, "Целевая архитектура с несколькими экземплярами")
    node_fill = "#EEF3FA"
    service_fill = "#F9F0DF"
    data_fill = "#EAF6EA"
    broker_fill = "#E9F7F7"
    draw_box(draw, (80, 170, 320, 280), "Клиент A", fill=node_fill)
    draw_box(draw, (80, 500, 320, 610), "Клиент B", fill=node_fill)
    draw_box(draw, (440, 330, 680, 450), "nginx", fill="#F3F3F3", bold=True)
    draw_box(draw, (820, 180, 1110, 300), "app-api-1", fill=service_fill, bold=True)
    draw_box(draw, (820, 500, 1110, 620), "app-api-2", fill=service_fill, bold=True)
    draw_box(draw, (1260, 330, 1510, 450), "Redis backplane", fill=broker_fill, bold=True)
    draw_box(draw, (1260, 650, 1510, 770), "PostgreSQL", fill=data_fill)
    draw_author_marker(draw, (1175, 270, 1560, 315), "Redis backplane: разработано")
    draw_author_marker(draw, (760, 95, 1180, 140), "SERVER_INSTANCE_ID + диагностика")
    draw_author_marker(draw, (760, 660, 1180, 705), "compose/nginx стенд")
    draw_arrow(draw, (320, 225), (440, 360))
    draw_arrow(draw, (320, 555), (440, 420))
    draw_arrow(draw, (680, 360), (820, 240))
    draw_arrow(draw, (680, 420), (820, 560))
    draw_arrow(draw, (1110, 240), (1260, 360), color="#0A6E73")
    draw_arrow(draw, (1260, 420), (1110, 560), color="#0A6E73")
    draw_arrow(draw, (1110, 270), (1260, 690), color="#4C7A38")
    draw_arrow(draw, (1110, 590), (1260, 730), color="#4C7A38")
    note_font = pil_font(24)
    note = "В рамках ВКР добавлены Redis backplane, диагностика и compose/nginx стенд с несколькими экземплярами."
    draw.rounded_rectangle((120, 790, 1480, 855), radius=14, fill="#FFF8D6", outline="#D2B43A", width=2)
    bbox = draw.textbbox((0, 0), note, font=note_font)
    draw.text(((1600 - (bbox[2] - bbox[0])) // 2, 812), note, fill="#111111", font=note_font)
    img.save(path)


def create_event_flow_figure(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    draw = ImageDraw.Draw(img)
    add_figure_title(draw, "Поток межузловой доставки события")
    labels = ["Клиент A", "app-api-1", "Redis backplane", "app-api-2", "Клиент B"]
    x_positions = [120, 450, 780, 1110, 1390]
    top = 150
    bottom = 760
    font = pil_font(24, bold=True)
    small = pil_font(21)
    for x, label in zip(x_positions, labels):
        draw_box(draw, (x - 105, 95, x + 105, 165), label, fill="#EEF3FA" if "Клиент" in label else "#F9F0DF")
        draw.line((x, top, x, bottom), fill="#999999", width=2)
    steps = [
        (0, 1, 230, "1. HTTP-команда изменения отчета"),
        (1, 1, 315, "2. Сохранение состояния в БД"),
        (1, 2, 400, "3. Публикация UI-события"),
        (2, 3, 485, "4. Доставка через Redis backplane"),
        (3, 4, 570, "5. Получение на другом узле"),
    ]
    for src, dst, y, label in steps:
        if src == dst:
            x = x_positions[src]
            draw.arc((x + 30, y - 25, x + 170, y + 45), start=90, end=270, fill="#111111", width=4)
            draw_arrow(draw, (x + 40, y + 35), (x + 35, y + 34), width=4)
            draw.text((x + 80, y - 20), label, fill="#111111", font=small)
        else:
            draw_arrow(draw, (x_positions[src] + 30, y), (x_positions[dst] - 30, y), color="#0A6E73")
            mid = (x_positions[src] + x_positions[dst]) // 2
            bbox = draw.textbbox((0, 0), label, font=small)
            draw.text((mid - (bbox[2] - bbox[0]) // 2, y - 35), label, fill="#111111", font=small)
    draw.rounded_rectangle((145, 795, 1455, 855), radius=14, fill="#FFF8D6", outline="#D2B43A", width=2)
    draw_author_marker(draw, (440, 650, 1160, 700), "ReportPageHubClient + проверочный Node.js-клиент")
    note = "Ключевой проверяемый факт: клиент B получает событие, хотя подключен к другому узлу."
    bbox = draw.textbbox((0, 0), note, font=small)
    draw.text(((1600 - (bbox[2] - bbox[0])) // 2, 814), note, fill="#111111", font=small)
    img.save(path)


def ensure_figures() -> list[Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    figures = [
        ASSET_DIR / "figure-1-single-node.png",
        ASSET_DIR / "figure-2-no-backplane.png",
        ASSET_DIR / "figure-3-scaleout.png",
        ASSET_DIR / "figure-4-event-flow.png",
    ]
    create_single_node_figure(figures[0])
    create_no_backplane_figure(figures[1])
    create_scaleout_figure(figures[2])
    create_event_flow_figure(figures[3])
    return figures


def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
            continue
        for key, value in edge_data.items():
            element.set(qn(f"w:{key}"), str(value))


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def clear_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(cell, top=None, left=None, bottom=None, right=None, insideH=None, insideV=None)


def set_bottom_border(cell):
    set_cell_borders(
        cell,
        top=None,
        left=None,
        right=None,
        insideH=None,
        insideV=None,
        bottom={"val": "single", "sz": "6", "space": "0", "color": "000000"},
    )


def set_page_margins(section):
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)


def set_run_font(run, *, size=14, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def setup_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.first_line_indent = Cm(1.25)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in [("Heading 1", 14), ("Heading 2", 14), ("Heading 3", 14)]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 6)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.page_break_before = False
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for name in ["Title", "Subtitle"]:
        try:
            style = styles[name]
        except KeyError:
            continue
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_page_number_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=12)


def enable_field_update_on_open(doc: Document):
    settings = doc.settings.element
    existing = settings.find(qn("w:updateFields"))
    if existing is None:
        existing = OxmlElement("w:updateFields")
        settings.append(existing)
    existing.set(qn("w:val"), "true")


def set_start_page_number(section, start: int):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def clear_start_page_number(section):
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is not None:
        sect_pr.remove(pg_num_type)


def add_centered(doc, text: str, *, size=14, bold=False, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def add_spacer(doc, points: int):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    return p


def add_left(doc, text: str, *, size=14, bold=False, first_line=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Cm(first_line)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    return p


def set_table_cell_margins(table, margin_twips: int = 0):
    for row in table.rows:
        for cell in row.cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_mar = tc_pr.first_child_found_in("w:tcMar")
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side in ("top", "left", "bottom", "right"):
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(margin_twips))
                node.set(qn("w:type"), "dxa")


def fill_template_cell(doc: Document, table_idx: int, row: int, col: int, text: str, *, size: int = 14, align=WD_ALIGN_PARAGRAPH.CENTER, bold: bool = False):
    cell = doc.tables[table_idx].rows[row].cells[col]
    set_cell_text(cell, text, size=size, bold=bold, align=align)


def _resize_existing_runs(cell, size_pt: int):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.size = Pt(size_pt)


_SHORT_TAIL_WORDS = {"с", "со", "в", "во", "на", "по", "для", "от", "до", "из", "к", "о", "об", "у", "и", "а", "или", "не"}


def split_into_lines(text: str, n_lines: int) -> list[str]:
    """Разбить строку по словам на n_lines примерно равных по длине частей.
    Избегаем переносов, где строка закончилась бы на коротком предлоге."""
    words = text.split()
    if not words:
        return [""] * n_lines
    total = len(text)
    target = total / n_lines
    lines: list[str] = []
    current: list[str] = []
    current_len = 0
    for word in words:
        new_len = current_len + (1 if current else 0) + len(word)
        last_word = current[-1].lower() if current else ""
        ends_with_preposition = last_word in _SHORT_TAIL_WORDS
        if (
            current
            and new_len > target
            and not ends_with_preposition
            and len(lines) < n_lines - 1
        ):
            lines.append(" ".join(current))
            current = [word]
            current_len = len(word)
        else:
            current.append(word)
            current_len = new_len
    if current:
        lines.append(" ".join(current))
    while len(lines) < n_lines:
        lines.append("")
    return lines[:n_lines]


def _freeze_paragraph_format(p):
    # Закрываем только те свойства, которые setup_styles меняет в стиле Normal.
    # space_before/space_after оставляем — они задают вертикальный воздух между
    # секциями титульного листа (БАКАЛАВРСКАЯ РАБОТА, Санкт-Петербург 2026, и т.п.).
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0


def _ensure_run_default_size(p, default_pt: int = 12):
    for r in p.runs:
        if r.font.size is None:
            r.font.size = Pt(default_pt)


def _strip_run_highlights(p):
    for r in p.runs:
        r_pr = r._r.find(qn("w:rPr"))
        if r_pr is None:
            continue
        for h in r_pr.findall(qn("w:highlight")):
            r_pr.remove(h)


def _drop_trailing_rows(table, *, keep: int):
    while len(table.rows) > keep:
        last = table.rows[-1]._tr
        last.getparent().remove(last)


def fill_lined_template_lines(
    doc: Document,
    table_idx: int,
    lines: list[str],
    *,
    size: int = 11,
):
    """Заполнить линованное поле заранее выбранными строками формы."""
    table = doc.tables[table_idx]
    for row_idx, line in enumerate(lines[: len(table.rows)]):
        col_idx = 1 if row_idx == 0 else 0
        cell = table.rows[row_idx].cells[col_idx]
        set_cell_text(cell, line, size=size, align=WD_ALIGN_PARAGRAPH.LEFT)


def _compact_break_paragraphs(doc: Document):
    """Параграфы-носители <w:br type=page/> или <w:sectPr/> ужимаем до
    точечной высоты, чтобы они не вытесняли титульник на отдельную страницу."""
    for p in doc.paragraphs:
        has_break = (
            p._p.find(".//" + qn("w:br") + "[@" + qn("w:type") + "='page']") is not None
            or p._p.find(".//" + qn("w:sectPr")) is not None
        )
        # Эмпирически: считаем "техническим" любой параграф без видимого текста
        # и без runs.
        has_text = any((t.text or "").strip() for t in p._p.iter(qn("w:t")))
        if has_break or not has_text:
            pPr = p._p.find(qn("w:pPr"))
            if pPr is None:
                pPr = OxmlElement("w:pPr")
                p._p.insert(0, pPr)
            existing = pPr.find(qn("w:spacing"))
            if existing is None:
                existing = OxmlElement("w:spacing")
                pPr.append(existing)
            existing.set(qn("w:line"), "20")
            existing.set(qn("w:lineRule"), "exact")
            existing.set(qn("w:before"), "0")
            existing.set(qn("w:after"), "0")


def neutralize_template_styles(doc: Document):
    for p in doc.paragraphs:
        _freeze_paragraph_format(p)
        _ensure_run_default_size(p)
        _strip_run_highlights(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _freeze_paragraph_format(p)
                    _ensure_run_default_size(p)
                    _strip_run_highlights(p)
    _compact_break_paragraphs(doc)


def load_vkr_title_pages() -> Document:
    doc = Document(str(VKR_TITLE_TEMPLATE))
    # Нейтрализуем наследование Normal-стиля и убираем подсветки до записи,
    # чтобы явные настройки от set_cell_text не были перезатёрты.
    neutralize_template_styles(doc)
    title_section = doc.sections[0]
    # Шаблон помечает свой sectPr как continuous, что после add_section превращается
    # в continuous-break и роняет первый абзац тела на ту же страницу. Убираем тип,
    # чтобы python-docx позже выставил nextPage через WD_SECTION.NEW_PAGE.
    sect_pr = title_section._sectPr
    for t in sect_pr.findall(qn("w:type")):
        sect_pr.remove(t)
    # Удаляем пустой параграф между последней таблицей и sectPr из шаблона —
    # он добавляется к section 0 и может перенести содержимое body на лишнюю страницу.
    body = doc.element.body
    children = list(body)
    for i, el in enumerate(children):
        if el.tag.endswith("}sectPr"):
            # Идём назад, удаляя пустые параграфы перед sectPr.
            j = i - 1
            while j >= 0 and children[j].tag.endswith("}p"):
                p = children[j]
                has_text = any((t.text or "").strip() for t in p.iter(qn("w:t")))
                if has_text:
                    break
                body.remove(p)
                j -= 1
            break

    # Титульный лист (страница 1).
    # Table 0 — заведующий кафедрой, уже заполнен в шаблоне; чистим только demo-дату подписи.
    fill_template_cell(doc, 0, 0, 2, "")

    # Table 1 — "на тему": тема ВКР распределяется по 2 строкам шаблона.
    # Третья строка остаётся пустой как декоративная (как в шаблоне).
    title_lines = split_into_lines(TITLE, 2)
    for r, line in enumerate(title_lines):
        fill_template_cell(doc, 1, r, 1, line, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Table 2 — "выполнена": ФИО в творительном падеже справа от метки.
    # Метку уменьшаем до 11pt, иначе «выполнена» не помещается в узкую ячейку 2.3 см.
    _resize_existing_runs(doc.tables[2].rows[0].cells[0], 11)
    fill_template_cell(doc, 2, 0, 1, STUDENT_INSTR, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Table 3 — направление подготовки уже заполнено в шаблоне (09.03.04 / Программная инженерия).
    # Направленность: код и название должны стоять в той же строке, что и метка.
    fill_template_cell(doc, 3, 4, 1, "02", size=14)
    fill_template_cell(doc, 3, 4, 3, "Проектирование программных систем", size=12)
    # В шаблоне название направленности продублировано в объединённой строке 6 — очищаем.
    for ci in range(4):
        fill_template_cell(doc, 3, 6, ci, "", size=12)

    # Table 4 — студент группы №: группа и фамилия. Чистим demo-дату.
    fill_template_cell(doc, 4, 0, 1, GROUP, size=14)
    fill_template_cell(doc, 4, 0, 3, "")
    fill_template_cell(doc, 4, 0, 5, STUDENT_SHORT, size=14)

    # Table 5 — Руководитель: должность, дата (очищаем), ФИО.
    fill_template_cell(doc, 5, 0, 0, VKR_SUPERVISOR_ROLE, size=12)
    fill_template_cell(doc, 5, 0, 2, "")
    fill_template_cell(doc, 5, 0, 4, VKR_SUPERVISOR, size=14)

    # Лист задания (страница 2).
    # Table 6 — заведующий кафедрой, уже заполнен; чистим demo-дату.
    fill_template_cell(doc, 6, 0, 2, "")

    # Table 7 — студенту группы N <ФИО>.
    fill_template_cell(doc, 7, 0, 1, GROUP, size=14)
    fill_template_cell(doc, 7, 0, 3, STUDENT_FULL, size=14)

    # Table 8 — тема задания: распределяется по 2 строкам, как и в Table 1.
    for r, line in enumerate(title_lines):
        fill_template_cell(doc, 8, r, 1, line, size=14, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Table 9 — приказ ГУАП: номер и дата оставляем пустыми, заполнятся вручную.

    # Tables 10-12 — линованные поля задания. Сохраняем строки шаблона и
    # заполняем только существующие ячейки, как в подписанных примерах.
    fill_lined_template_lines(doc, 10, VKR_GOAL_LINES)
    fill_lined_template_lines(doc, 11, VKR_TASK_LINES)
    fill_lined_template_lines(doc, 12, VKR_CONTENT_LINES)

    # Table 13 — срок сдачи: чистим demo-дату «15 июня 2026», подпишут от руки.
    for ci in (1, 3, 5):
        fill_template_cell(doc, 13, 0, ci, "")

    # Table 14 — Руководитель: должность, дата (очищаем), ФИО.
    fill_template_cell(doc, 14, 0, 0, VKR_SUPERVISOR_ROLE, size=12)
    fill_template_cell(doc, 14, 0, 2, "")
    fill_template_cell(doc, 14, 0, 4, VKR_SUPERVISOR, size=14)

    # Table 15 — Задание принял к исполнению студент группы.
    fill_template_cell(doc, 15, 0, 1, GROUP, size=14)
    fill_template_cell(doc, 15, 0, 3, "")
    fill_template_cell(doc, 15, 0, 5, STUDENT_SHORT, size=14)

    return doc


def add_practice_title_page(doc: Document):
    section = doc.sections[-1]
    section.top_margin = Cm(1.0)
    section.bottom_margin = Cm(0.8)

    add_centered(doc, "МИНИСТЕРСТВО ОБРАЗОВАНИЯ И НАУКИ РОССИЙСКОЙ ФЕДЕРАЦИИ", size=11)
    add_centered(doc, "федеральное государственное автономное образовательное учреждение высшего образования", size=10)
    add_centered(doc, "«Санкт–Петербургский государственный университет", size=11)
    add_centered(doc, "аэрокосмического приборостроения»", size=11, space_after=16)
    add_centered(doc, "Кафедра №43 «Компьютерных технологий и программной инженерии»", size=11, space_after=16)

    add_left(doc, "ОТЧЁТ ПО ПРАКТИКЕ", size=12, bold=True)
    add_left(doc, "ЗАЩИЩЁН С ОЦЕНКОЙ", size=12)
    add_left(doc, "РУКОВОДИТЕЛЬ", size=12, first_line=0, space_after=2)

    supervisor = doc.add_table(rows=2, cols=3)
    clear_table_borders(supervisor)
    set_table_cell_margins(supervisor, 0)
    for row in supervisor.rows:
        row.cells[0].width = Cm(5.6)
        row.cells[1].width = Cm(3.8)
        row.cells[2].width = Cm(4.8)
    set_cell_text(supervisor.cell(0, 0), PRACTICE_SUPERVISOR_ROLE, size=11)
    set_cell_text(supervisor.cell(0, 1), "", size=11)
    set_cell_text(supervisor.cell(0, 2), PRACTICE_SUPERVISOR, size=11)
    for c in range(3):
        set_bottom_border(supervisor.cell(0, c))
    set_cell_text(supervisor.cell(1, 0), "должность, уч. степень, звание", size=7)
    set_cell_text(supervisor.cell(1, 1), "подпись, дата", size=7)
    set_cell_text(supervisor.cell(1, 2), "инициалы, фамилия", size=7)

    add_spacer(doc, 12)
    add_centered(doc, "ОТЧЁТ ПО ПРАКТИКЕ", size=14, bold=True, space_after=4)

    meta = doc.add_table(rows=14, cols=3)
    clear_table_borders(meta)
    set_table_cell_margins(meta, 0)
    for row in meta.rows:
        row.cells[0].width = Cm(5.2)
        row.cells[1].width = Cm(5.2)
        row.cells[2].width = Cm(5.2)
    rows = [
        ("вид практики", PRACTICE_TYPE, ""),
        ("", "вид практики", ""),
        ("тип практики", PRACTICE_KIND, ""),
        ("", "тип практики", ""),
        ("на тему индивидуального задания", TITLE, ""),
        ("", "на тему индивидуального задания", ""),
        ("выполнен", STUDENT_INSTR, ""),
        ("", "фамилия, имя, отчество обучающегося в творительном падеже", ""),
        ("по направлению подготовки", "09.03.04", "Программная инженерия"),
        ("", "код", "наименование направления"),
        ("направленности", "02", "Проектирование программных систем"),
        ("", "код", "наименование направленности"),
        ("Обучающийся группы №", GROUP, "____________________"),
        ("", "номер", "подпись, дата"),
    ]
    for r, values in enumerate(rows):
        if r == 4:
            meta.cell(r, 1).merge(meta.cell(r, 2))
        for c, value in enumerate(values):
            if r == 4 and c == 2:
                continue
            is_hint = r in {1, 3, 5, 7, 9, 11, 13}
            size = 7 if is_hint else 10
            set_cell_text(meta.cell(r, c), value, size=size, align=WD_ALIGN_PARAGRAPH.CENTER)
            if r in {0, 2, 4, 6, 8, 10, 12} and c > 0:
                set_bottom_border(meta.cell(r, c))

    sig = doc.add_table(rows=2, cols=3)
    clear_table_borders(sig)
    set_table_cell_margins(sig, 0)
    for row in sig.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(5)
        row.cells[2].width = Cm(5)
    set_cell_text(sig.cell(0, 0), "", size=7)
    set_cell_text(sig.cell(0, 1), "", size=7)
    set_cell_text(sig.cell(0, 2), "инициалы, фамилия", size=7)
    set_cell_text(sig.cell(1, 0), "", size=10)
    set_cell_text(sig.cell(1, 1), "", size=10)
    set_cell_text(sig.cell(1, 2), STUDENT_SHORT, size=10)
    set_bottom_border(sig.cell(1, 2))

    add_spacer(doc, 20)
    add_centered(doc, "Санкт–Петербург 2026", size=10)
    doc.add_page_break()

    set_page_margins(section)


def add_field(paragraph, instr_text: str):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Обновите поле в Word"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)
    set_run_font(run)


def add_abstract(doc: Document):
    p = doc.add_heading("РЕФЕРАТ", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    facts = (
        "Выпускная квалификационная работа содержит 51 страницу, 6 рисунков, "
        "2 таблицы, список использованных источников из 20 наименований и 3 приложения."
    )
    add_body_paragraph(doc, facts)
    add_body_paragraph(
        doc,
        "Ключевые слова: real-time интерфейс, WebSocket, SignalR, горизонтальное масштабирование, Redis backplane, система отслеживания задач, межузловая доставка событий.",
    )
    add_body_paragraph(
        doc,
        "Актуальность работы обусловлена необходимостью обеспечения корректного real-time взаимодействия пользователей системы отслеживания задач при переходе от одного экземпляра серверного приложения к многосерверной конфигурации.",
    )
    add_body_paragraph(
        doc,
        "Цель работы — разработать и исследовать real-time интерфейс для системы отслеживания задач, в котором WebSocket-соединения обслуживаются несколькими экземплярами серверного приложения, а события об изменении сущностей корректно доставляются клиентам независимо от узла подключения.",
    )
    add_body_paragraph(
        doc,
        "Полученные результаты: разработан распределенный real-time контур на базе SignalR и Redis backplane, подготовлен воспроизводимый стенд с несколькими экземплярами, реализованы диагностические средства и проведена экспериментальная проверка межузловой доставки событий.",
    )
    doc.add_page_break()


def add_toc_line(doc: Document, title: str, page: int, *, level: int = 1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(0 if level == 1 else 0.75)
    p.paragraph_format.right_indent = Cm(0)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.0), WD_TAB_LEADER.DOTS)
    run = p.add_run(f"{title}\t{page}")
    set_run_font(run, size=11, bold=level == 1)


def add_toc(doc: Document, *, kind: str = "vkr"):
    p = doc.add_heading("СОДЕРЖАНИЕ", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if kind == "practice":
        entries = [
            (1, "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ", 3),
            (1, "ВВЕДЕНИЕ", 4),
            (1, "1. Анализ предметной области и существующих подходов", 7),
            (2, "1.1 Описание использованной литературы", 7),
            (2, "1.2 Real-time взаимодействие в системах отслеживания задач", 9),
            (2, "1.3 WebSocket и SignalR как основа real-time контура", 10),
            (2, "1.4 Ограничение single-node WebSocket-архитектуры", 11),
            (2, "1.5 Сравнение подходов к масштабированию", 12),
            (2, "1.6 Вывод по главе", 12),
            (1, "2. Анализ целевой системы и постановка задачи практики", 14),
            (2, "2.1 Характеристика целевой системы", 14),
            (2, "2.2 Исходная реализация real-time взаимодействия", 15),
            (2, "2.3 Требования к разрабатываемому решению", 15),
            (2, "2.4 Выбор платформы и инструментальных средств", 16),
            (2, "2.5 Постановка задачи", 16),
            (1, "3. Проектирование масштабируемого real-time контура", 18),
            (2, "3.1 Исходная архитектура", 18),
            (2, "3.2 Целевая архитектура", 18),
            (2, "3.3 Поток real-time события", 19),
            (2, "3.4 Модель групп и повторной подписки", 20),
            (2, "3.5 Диагностический контур", 21),
            (1, "4. Выполнение индивидуального задания", 22),
            (2, "4.1 Организация работ", 22),
            (2, "4.2 Изменения backend", 22),
            (2, "4.3 Проверочный клиент", 23),
            (2, "4.4 Инфраструктурные изменения", 23),
            (2, "4.5 Автоматизированный сценарий проверки", 24),
            (1, "5. Результаты практики и оценка решения", 26),
            (2, "5.1 Цель и программа испытаний", 26),
            (2, "5.2 Стенд испытаний", 26),
            (2, "5.3 Результат без Redis backplane", 27),
            (2, "5.4 Результат с Redis backplane", 27),
            (2, "5.5 Серийная проверка и задержка доставки", 28),
            (2, "5.6 Проверка восстановления соединения и отказа узла", 29),
            (2, "5.7 Сравнительная таблица результатов", 29),
            (2, "5.8 Проверка входной точки стенда", 30),
            (2, "5.9 Оценка достоверности результатов", 30),
            (2, "5.10 Вывод по испытаниям", 31),
            (1, "ЗАКЛЮЧЕНИЕ", 31),
            (1, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 33),
            (1, "ПРИЛОЖЕНИЕ А", 34),
        ]
    else:
        entries = [
            (1, "РЕФЕРАТ", 3),
            (1, "ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ", 5),
            (1, "ВВЕДЕНИЕ", 6),
            (1, "1. Анализ предметной области и существующих подходов", 10),
            (2, "1.1 Real-time взаимодействие в системах отслеживания задач", 10),
            (2, "1.2 WebSocket и SignalR как основа real-time контура", 10),
            (2, "1.3 Ограничение одноузловой WebSocket-архитектуры", 11),
            (2, "1.4 Сравнение подходов к масштабированию", 12),
            (2, "1.5 Вывод по главе", 14),
            (1, "2. Анализ целевой системы и постановка задачи", 15),
            (2, "2.1 Характеристика целевой системы", 15),
            (2, "2.2 Исходная реализация real-time взаимодействия", 16),
            (2, "2.3 Требования к разрабатываемому решению", 17),
            (2, "2.4 Выбор платформы и инструментальных средств", 18),
            (2, "2.5 Постановка задачи", 18),
            (1, "3. Проектирование масштабируемого real-time контура", 19),
            (2, "3.1 Исходная архитектура", 19),
            (2, "3.2 Целевая архитектура", 20),
            (2, "3.3 Поток real-time события", 21),
            (2, "3.4 Модель групп и повторной подписки", 22),
            (2, "3.5 Диагностический контур", 22),
            (2, "3.6 Семантика доставки и ограничения", 23),
            (1, "4. Реализация решения в целевой системе", 25),
            (2, "4.1 Организация работ", 25),
            (2, "4.2 Изменения серверной части", 25),
            (2, "4.3 Изменения клиентской части", 27),
            (2, "4.4 Проверочный клиент", 28),
            (2, "4.5 Инфраструктурные изменения", 28),
            (2, "4.6 Автоматизированный сценарий проверки", 29),
            (1, "5. Испытания и оценка результатов", 31),
            (2, "5.1 Цель и программа испытаний", 31),
            (2, "5.2 Стенд испытаний", 31),
            (2, "5.3 Результат без Redis backplane", 32),
            (2, "5.4 Результат с Redis backplane", 32),
            (2, "5.5 Серийная проверка и задержка доставки", 33),
            (2, "5.6 Проверка восстановления соединения и отказа узла", 34),
            (2, "5.7 Сравнительная таблица результатов", 34),
            (2, "5.8 Проверка входной точки стенда", 35),
            (2, "5.9 Оценка достоверности результатов", 35),
            (2, "5.10 Вывод по испытаниям", 36),
            (1, "ЗАКЛЮЧЕНИЕ", 37),
            (1, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", 39),
            (1, "ПРИЛОЖЕНИЕ А", 40),
            (1, "ПРИЛОЖЕНИЕ Б", 41),
            (1, "ПРИЛОЖЕНИЕ В", 45),
        ]
    for level, title, page in entries:
        add_toc_line(doc, title, page, level=level)
    doc.add_page_break()


def add_abbreviations(doc: Document, *, kind: str = "vkr"):
    p = doc.add_heading("ПЕРЕЧЕНЬ СОКРАЩЕНИЙ И ОБОЗНАЧЕНИЙ", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data = [
        ("API", "Application Programming Interface, программный интерфейс приложения"),
        ("ASP.NET Core", "платформа для разработки веб-приложений на языке C#"),
        ("HTTP", "HyperText Transfer Protocol, протокол передачи гипертекста"),
        ("JSON", "JavaScript Object Notation, текстовый формат обмена данными"),
        ("nginx", "обратный прокси-сервер и балансировщик нагрузки"),
        ("Redis", "система хранения данных в памяти, используемая как backplane"),
        ("SignalR", "библиотека ASP.NET Core для real-time взаимодействия"),
        ("UI", "User Interface, пользовательский интерфейс"),
        ("URL", "Uniform Resource Locator, адрес ресурса"),
        ("WebSocket", "сетевой протокол двусторонней связи поверх одного TCP-соединения"),
    ]
    if kind == "vkr":
        data.append(("ВКР", "выпускная квалификационная работа"))
    else:
        data.extend(
            [
                ("БД", "база данных"),
                ("ИЗ", "индивидуальное задание"),
                ("ПО", "программное обеспечение"),
            ]
        )
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(4)
    table.columns[1].width = Cm(12)
    set_cell_text(table.cell(0, 0), "Сокращение", bold=True)
    set_cell_text(table.cell(0, 1), "Расшифровка", bold=True)
    for abbr, meaning in data:
        cells = table.add_row().cells
        set_cell_text(cells[0], abbr, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], meaning, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph(style="Normal")
    has_inline_code = "`" in text
    has_long_token = any(len(token) > 28 for token in re.findall(r"\S+", text))
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if has_inline_code or has_long_token else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.line_spacing = 1.5
    add_inline_runs(p, text)
    return p


def add_list_item(doc: Document, text: str, *, numbered: bool = True, index: int = 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    prefix = f"{index}. " if numbered else "— "
    add_inline_runs(p, prefix + text)
    return p


def add_figure(doc: Document, image_path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(15.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.space_after = Pt(6)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=12, italic=True)


def add_code_block(doc: Document, code: str, language: str | None, figure_paths: list[Path], figure_idx: int) -> int:
    if language == "mermaid":
        captions = [
            "Рисунок 3 — Исходная одноузловая архитектура real-time контура",
            "Рисунок 4 — Архитектура с несколькими экземплярами без Redis backplane и фрагментация локальных SignalR-групп",
            "Рисунок 5 — Целевая архитектура с несколькими экземплярами, Redis backplane и элементами, разработанными в рамках ВКР",
            "Рисунок 6 — Последовательность межузловой доставки события ReceiveReportPatch",
        ]
        if figure_idx < len(figure_paths):
            add_figure(doc, figure_paths[figure_idx], captions[figure_idx])
        return figure_idx + 1
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_cell_margins(table, 90)
    cell = table.cell(0, 0)
    cell.width = Cm(15.5)
    set_cell_shading(cell, "F7F7F7")
    cell.text = ""
    lines = code.strip("\n").splitlines() or [""]
    first = True
    for original_line in lines:
        wrapped = textwrap.wrap(
            original_line,
            width=92,
            subsequent_indent="    ",
            replace_whitespace=False,
            drop_whitespace=False,
            break_long_words=True,
            break_on_hyphens=False,
        ) or [""]
        for line in wrapped:
            p = cell.paragraphs[0] if first else cell.add_paragraph()
            first = False
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.left_indent = Cm(0)
            p.paragraph_format.line_spacing = 0.95
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            set_run_font(run, size=8, name="Courier New")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return figure_idx


def set_row_as_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is None:
        existing = OxmlElement("w:tblHeader")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:cantSplit"))
    if existing is None:
        existing = OxmlElement("w:cantSplit")
        tr_pr.append(existing)
    existing.set(qn("w:val"), "true")


def add_markdown_table(doc: Document, lines: list[str], table_idx: int) -> int:
    rows = []
    for line in lines:
        stripped = line.strip()
        if not stripped.startswith("|"):
            continue
        cells = [cell.strip() for cell in stripped.strip("|").split("|")]
        if all(set(cell) <= {"-", ":", " "} for cell in cells):
            continue
        rows.append(cells)
    if not rows:
        return table_idx

    captions = {
        "N": "Сравнение подходов к масштабированию WebSocket-соединений",
        "Подход": "Сравнение подходов к масштабированию WebSocket-соединений",
        "Режим": "Сравнительная таблица результатов испытаний",
    }
    caption_text = captions.get(rows[0][0], "Сравнительные данные")
    is_results_table = rows[0][0] == "Режим"
    is_approach_table = rows[0][0] == "N"

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cap.paragraph_format.first_line_indent = Cm(0)
    cap.paragraph_format.keep_with_next = True
    cap.paragraph_format.space_after = Pt(2)
    run = cap.add_run(f"Таблица {table_idx} — {caption_text}")
    set_run_font(run, size=12, bold=True)

    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    cell_size = 9 if is_approach_table else (10 if is_results_table else 11)
    for c, text in enumerate(rows[0]):
        set_cell_text(table.cell(0, c), text, size=cell_size, bold=True)
    set_row_as_header(table.rows[0])
    set_row_cant_split(table.rows[0])
    for row in rows[1:]:
        row_obj = table.add_row()
        cells = row_obj.cells
        for c, text in enumerate(row[: len(cells)]):
            set_cell_text(cells[c], text, size=cell_size, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_row_cant_split(row_obj)

    if is_approach_table or is_results_table:
        usable_width = Cm(16.5)
        weights = ([0.6, 2.6, 2.1, 4.4, 3.8] if is_approach_table else [2.4, 1.0, 2.6, 2.6, 2.2])[: len(rows[0])]
        total = sum(weights)
        col_widths = [Cm(usable_width.cm * w / total) for w in weights]
        for c, w in enumerate(col_widths):
            table.columns[c].width = w
            for row in table.rows:
                row.cells[c].width = w

    return table_idx + 1


def add_sources_as_numbered_list(doc: Document, source_lines: list[str]):
    p = doc.add_heading("СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for index, line in enumerate(source_lines, start=1):
        m = re.match(r"\d+\.\s+(.*)", line.strip())
        if not m:
            continue
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(f"{index}. {m.group(1)}")
        set_run_font(run, size=12)


def transform_for_practice(text: str) -> str:
    replacements = {
        "## 2. Анализ целевой системы и постановка задачи": "## 2. Анализ целевой системы и постановка задачи практики",
        "## 4. Реализация решения в целевой системе": "## 4. Выполнение индивидуального задания",
        "## 5. Испытания и оценка результатов": "## 5. Результаты практики и оценка решения",
        "Цель работы: разработать и исследовать real-time интерфейс": "Цель преддипломной практики: разработать и исследовать real-time интерфейс",
        "Для достижения цели поставлены следующие задачи:": "В ходе практики решались следующие задачи:",
        "Элемент новизны работы заключается в инженерно-исследовательском обосновании и экспериментальной проверке распределенного": "Результатом практики является инженерно-исследовательское обоснование и экспериментальная проверка распределенного",
        "Новизна работы заключается": "Результатом практики является",
        "Элемент новизны работы заключается": "Результатом практики является",
        "В рамках работы не разрабатывается новый сетевой протокол, но выполняется авторская композиция": "В рамках практики не разрабатывался новый сетевой протокол, но была выполнена авторская композиция",
        "Практическая значимость работы состоит": "Практическая значимость результатов практики состоит",
        "В рамках ВКР": "В рамках преддипломной практики",
        "Для цели данной ВКР": "Для цели преддипломной практики",
        "Для ВКР подготовлен отдельный compose-стенд": "Для практики подготовлен отдельный compose-стенд",
        "В ходе работы была рассмотрена": "В ходе преддипломной практики была рассмотрена",
        "В отдельной ветке проекта реализована": "При выполнении индивидуального задания в отдельной ветке проекта реализована",
        "Полученные результаты имеют практическую значимость": "Полученные в ходе практики результаты имеют практическую значимость",
        "заявленный тезис работы": "заявленный результат практики",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace(
        "## 1. Анализ предметной области и существующих подходов\n\n",
        "## 1. Анализ предметной области и существующих подходов\n\n"
        + PRACTICE_LITERATURE_REVIEW.strip()
        + "\n\n",
    )
    practice_heading_replacements = {
        "### 1.1 Real-time взаимодействие в системах отслеживания задач": "### 1.2 Real-time взаимодействие в системах отслеживания задач",
        "### 1.2 WebSocket и SignalR как основа real-time контура": "### 1.3 WebSocket и SignalR как основа real-time контура",
        "### 1.3 Ограничение single-node WebSocket-архитектуры": "### 1.4 Ограничение single-node WebSocket-архитектуры",
        "### 1.4 Сравнение подходов к масштабированию": "### 1.5 Сравнение подходов к масштабированию",
        "### 1.5 Вывод по главе": "### 1.6 Вывод по главе",
    }
    for old, new in practice_heading_replacements.items():
        text = text.replace(old, new)
    text = text.replace(
        "Научно-техническая проблема работы состоит",
        "Основная задача преддипломной практики состоит",
    )
    text = text.replace(
        "Объектом исследования являются",
        "Объектом рассмотрения в отчете являются",
    )
    text = text.replace(
        "Предметом исследования являются",
        "Предметом практической работы являются",
    )
    return text


def process_markdown(doc: Document, text: str, figure_paths: list[Path], *, kind: str = "vkr"):
    lines = text.splitlines()
    in_code = False
    code_lang = None
    code_lines: list[str] = []
    table_lines: list[str] = []
    table_idx = 1
    figure_idx = 0
    skip_until_main = True
    in_sources = False
    source_lines: list[str] = []
    numbered_counter = 0
    last_was_numbered = False

    def flush_table():
        nonlocal table_lines, table_idx
        if table_lines:
            table_idx = add_markdown_table(doc, table_lines, table_idx)
            table_lines = []

    def reset_numbering():
        nonlocal numbered_counter, last_was_numbered
        numbered_counter = 0
        last_was_numbered = False

    for raw in lines:
        line = raw.rstrip()

        if skip_until_main:
            if line.strip() == "## Введение":
                skip_until_main = False
            else:
                continue

        if line.startswith("## Список рисунков") or line.startswith("## Технические приложения") or line.startswith("## Что нужно сделать"):
            break

        if in_code:
            if line.startswith("```"):
                figure_idx = add_code_block(doc, "\n".join(code_lines), code_lang, figure_paths, figure_idx)
                in_code = False
                code_lang = None
                code_lines = []
            else:
                code_lines.append(line)
            continue

        if line.startswith("```"):
            flush_table()
            in_code = True
            code_lang = line.strip("`").strip() or None
            code_lines = []
            reset_numbering()
            continue

        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)$", line.strip())
        if image_match:
            caption = image_match.group(1)
            image_ref = Path(image_match.group(2))
            image_path = image_ref if image_ref.is_absolute() else ROOT / image_ref
            add_figure(doc, image_path, caption)
            reset_numbering()
            continue

        if line.strip().startswith("|"):
            table_lines.append(line)
            reset_numbering()
            continue
        flush_table()

        if not line.strip():
            reset_numbering()
            continue

        if line.strip() == "<!-- PAGE_BREAK -->":
            doc.add_page_break()
            reset_numbering()
            continue

        if line.strip() == "## Список использованных источников":
            in_sources = True
            reset_numbering()
            continue

        if in_sources:
            if re.match(r"\d+\.\s+", line.strip()):
                source_lines.append(line.strip())
            continue

        if line.startswith("## "):
            heading = line[3:].strip()
            p = doc.add_heading(heading.upper() if heading in {"Введение", "Заключение"} else heading, level=1)
            if heading in {"Введение", "Заключение"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if kind == "vkr" and heading == "Заключение":
                    p.paragraph_format.page_break_before = True
            else:
                p.paragraph_format.page_break_before = True
            reset_numbering()
            continue

        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            reset_numbering()
            continue

        numbered_match = re.match(r"^(\d+)\.\s+(.*)", line.strip())
        bullet_match = re.match(r"^[-*]\s+(.*)", line.strip())
        if numbered_match:
            if not last_was_numbered:
                numbered_counter = 1
            else:
                numbered_counter += 1
            add_list_item(doc, numbered_match.group(2), numbered=True, index=numbered_counter)
            last_was_numbered = True
        elif bullet_match:
            add_list_item(doc, bullet_match.group(1), numbered=False)
            last_was_numbered = False
        else:
            add_body_paragraph(doc, line)
            last_was_numbered = False

    flush_table()
    if source_lines:
        doc.add_page_break()
        add_sources_as_numbered_list(doc, source_lines)


def add_appendices(doc: Document, *, kind: str = "vkr"):
    p = doc.add_heading("ПРИЛОЖЕНИЕ А", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    add_centered(doc, "Конфигурация стенда", bold=True)
    if kind == "practice":
        add_body_paragraph(
            doc,
            "Ниже приведены ключевые элементы стенда, подтверждающие выполнение индивидуального задания: запуск двух экземпляров backend-сервиса, подключение Redis backplane и режим воспроизведения исходного ограничения без межузловой доставки.",
        )
        add_code_block(
            doc,
            """app-api-1:
  environment:
    REDIS_CONNECTION_STRING: redis:6379
    SERVER_INSTANCE_ID: app-api-1

app-api-2:
  environment:
    REDIS_CONNECTION_STRING: redis:6379
    SERVER_INSTANCE_ID: app-api-2

upstream app-api {
    server app-api-1:7777;
    server app-api-2:7777;
}""",
            "yaml",
            [],
            0,
        )
        add_body_paragraph(
            doc,
            "Для режима без backplane используется тот же стенд, но переменная `REDIS_CONNECTION_STRING` задается пустой строкой для обоих экземпляров `app-api`.",
        )
        add_code_block(
            doc,
            """Multi-instance без backplane:
  Error: ReceiveReportPatch timed out after 10000ms

Multi-instance с Redis backplane:
  ok: true
  nodeA.serverInstanceId: app-api-1
  nodeB.serverInstanceId: app-api-2
  deliveryLatencyMs: 9.5

Серийная проверка:
  iterations: 30
  successful: 30
  failed: 0
  avg: 7.2 ms
  p95: 13.8 ms""",
            "text",
            [],
            0,
        )
        return

    add_body_paragraph(
        doc,
        "Фрагмент `docker-compose.thesis.yml` задает два экземпляра серверного сервиса, общий Redis и диагностические идентификаторы серверных узлов.",
    )
    add_code_block(
        doc,
        """app-api-1:
  build:
    context: ./backend/bugget-api
    dockerfile: Bugget/Dockerfile
  environment:
    POSTGRES_CONNECTION_STRING: Host=db;Port=5432;Database=app_db;Username=postgres
    REDIS_CONNECTION_STRING: redis:6379
    SERVER_INSTANCE_ID: app-api-1
  ports: ["7771:7777"]

app-api-2:
  build:
    context: ./backend/bugget-api
    dockerfile: Bugget/Dockerfile
  environment:
    POSTGRES_CONNECTION_STRING: Host=db;Port=5432;Database=app_db;Username=postgres
    REDIS_CONNECTION_STRING: redis:6379
    SERVER_INSTANCE_ID: app-api-2
  ports: ["7772:7777"]

redis:
  image: redis:8
  container_name: redis_app_thesis""",
        "yaml",
        [],
        0,
    )
    add_body_paragraph(doc, "Фрагмент `docker-compose.thesis.no-backplane.yml` используется для воспроизведения исходного ограничения режима с несколькими экземплярами.")
    add_code_block(
        doc,
        """services:
  app-api-1:
    environment:
      REDIS_CONNECTION_STRING: ""

  app-api-2:
    environment:
      REDIS_CONNECTION_STRING: """"",
        "yaml",
        [],
        0,
    )
    add_body_paragraph(doc, "Фрагмент `upstreams.thesis.conf` задает upstream-блок nginx для балансировки HTTP- и WebSocket-трафика между двумя экземплярами `app-api`.")
    add_code_block(
        doc,
        """upstream app-api {
    zone app-api 64k;
    server app-api-1:7777;
    server app-api-2:7777;
}""",
        "nginx",
        [],
        0,
    )

    p = doc.add_heading("ПРИЛОЖЕНИЕ Б", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    add_centered(doc, "Фрагменты реализации серверной части", bold=True)
    add_body_paragraph(doc, "Настройка SignalR-конвейера выполнена в методе `AddMessaging` расширения `ServiceCollectionExtensions`. Сначала создается базовая регистрация хаба и JSON-протокола, затем условно подключается Redis backplane: при наличии переменной окружения `REDIS_CONNECTION_STRING` тот же исполняемый код переключается в режим межузловой доставки, при пустом значении — продолжает работать в одноузловом режиме. Это позволяет запускать один и тот же образ в позитивном и негативном контрольных режимах эксперимента.")
    add_code_block(
        doc,
        """public static IServiceCollection AddMessaging(this IServiceCollection services)
{
    services.AddSingleton<IUserIdProvider, SignalRUserIdProvider>();

    var signalRBuilder = services.AddSignalR(options =>
    {
        options.EnableDetailedErrors = true;
        options.KeepAliveInterval = TimeSpan.FromSeconds(15);
        options.ClientTimeoutInterval = TimeSpan.FromSeconds(60);
    })
    .AddJsonProtocol(options =>
    {
        options.PayloadSerializerOptions.DefaultIgnoreCondition =
            JsonIgnoreCondition.WhenWritingNull;
    })
    .AddHubOptions<ReportPageHub>(options =>
    {
        options.AddFilter<HubExceptionHandlerFilter>();
    });

    var redisConnectionString =
        Environment.GetEnvironmentVariable(EnvironmentConstants.RedisConnectionString);
    if (!string.IsNullOrWhiteSpace(redisConnectionString))
    {
        signalRBuilder.AddStackExchangeRedis(redisConnectionString, options =>
        {
            options.Configuration.ChannelPrefix =
                RedisChannel.Literal("app-api-realtime");
        });
    }

    return services;
}

// Регистрация централизованного отправителя событий в группу
services.AddSingleton<IReportPageHubClient, ReportPageHubClient>();
// Информация об экземпляре сервера, инжектируется в хаб и в отправитель событий
services.AddSingleton<ServerInstanceInfo>();""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "Контракт диагностических данных оформлен как `record` и используется в качестве возвращаемого типа метода хаба `GetConnectionDiagnosticsAsync`. В контракт включены идентификатор экземпляра сервера, имя машины, идентификатор SignalR-соединения и пользовательский идентификатор.")
    add_code_block(
        doc,
        """public sealed record RealtimeConnectionDiagnostics(
    string ServerInstanceId,
    string MachineName,
    string ConnectionId,
    string? UserIdentifier);""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "Диагностический метод хаба возвращает идентификатор серверного экземпляра и соединения, что позволяет доказать подключение клиентов к разным узлам и сопоставить серверные логи с действиями конкретного клиента.")
    add_code_block(
        doc,
        """public Task<RealtimeConnectionDiagnostics> GetConnectionDiagnosticsAsync()
{
    return Task.FromResult(new RealtimeConnectionDiagnostics(
        serverInstanceInfo.Id,
        serverInstanceInfo.MachineName,
        Context.ConnectionId,
        Context.UserIdentifier
    ));
}""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "Жизненный цикл соединения сопровождается журналированием с указанием идентификатора экземпляра сервера. Это позволяет в негативном и позитивном сценариях соотнести события подключения и отключения клиента с конкретным узлом.")
    add_code_block(
        doc,
        """public override async Task OnConnectedAsync()
{
    logger.LogWarning(
        "Клиент {@ConnectionId} подключился к SignalR на экземпляре {@ServerInstanceId}",
        Context.ConnectionId,
        serverInstanceInfo.Id);

    await base.OnConnectedAsync();
}

public override async Task OnDisconnectedAsync(Exception? exception)
{
    logger.LogWarning(
        "Клиент {@ConnectionId} отключился от экземпляра {@ServerInstanceId}. Причина: {@Reason}",
        Context.ConnectionId,
        serverInstanceInfo.Id,
        exception?.Message ?? "неизвестно");

    await base.OnDisconnectedAsync(exception);
}""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "Подписка клиента на real-time события отчета реализована как доменная операция, а не как прямой вызов `Groups.AddToGroupAsync`. Метод проверяет авторизацию пользователя, разрешает идентификатор отчета по нескольким источникам (внутренний, публичный и командный) и формирует ключ группы через структуру `ReportIdContext`. Только после доменной проверки выполняется присоединение соединения к группе SignalR.")
    add_code_block(
        doc,
        """public async Task JoinReportGroupAsync(string aliasId)
{
    var user = Context.User?.GetIdentity();
    if (user is null)
    {
        throw new HubException("пользователь не авторизован");
    }

    var (reportId, publicId, teamReportId) =
        ReportIdResolveHelper.ResolveReportId(aliasId, aliasOptions.Value);
    var resolvedReport = await reportsService.ResolveReportIdAsync(
        user.OrganizationId,
        user.TeamId,
        reportId,
        publicId,
        teamReportId);

    if (resolvedReport == null)
    {
        throw new HubException("репорт не найден");
    }

    var groupKey = new ReportIdContext(
        resolvedReport.Id,
        aliasId,
        resolvedReport.CreatorTeamId
    ).GroupKey;

    logger.LogWarning(
        "Клиент {@ConnectionId} подключился к группе {@ReportId} на экземпляре {@ServerInstanceId}",
        Context.ConnectionId,
        groupKey,
        serverInstanceInfo.Id);

    await Groups.AddToGroupAsync(Context.ConnectionId, groupKey);
}""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "Централизованная отправка события в группу сохраняет существующее содержимое событий и добавляет диагностическое логирование. Поддерживается режим исключения соединения-инициатора через `GroupExcept`, что позволяет не дублировать клиенту-инициатору событие, на которое он уже среагировал по HTTP-ответу.")
    add_code_block(
        doc,
        """private Task SendToGroupAsync(
    string groupKey,
    string eventName,
    object?[] args,
    string? excludedConnectionId = null)
{
    var eventId = Guid.NewGuid().ToString("N");

    logger.LogInformation(
        "Realtime event {@EventId} {@EventName} отправлен из {@ServerInstanceId} в группу {@GroupKey}, excludedConnectionId={@ExcludedConnectionId}",
        eventId,
        eventName,
        serverInstanceInfo.Id,
        groupKey,
        excludedConnectionId);

    var clients = excludedConnectionId is null
        ? hubContext.Clients.Group(groupKey)
        : hubContext.Clients.GroupExcept(groupKey, excludedConnectionId);

    return clients.SendCoreAsync(eventName, args);
}""",
        "csharp",
        [],
        0,
    )
    add_body_paragraph(doc, "На основе единого метода `SendToGroupAsync` построен набор типизированных доменных отправителей. Real-time контур покрывает не одно событие, а полный набор сущностей страницы отчета: сам отчет, баги, шаги воспроизведения, комментарии, вложения и ссылки. Ниже приведена выжимка из 17 методов класса `ReportPageHubClient`, показывающая разнообразие транслируемых событий и переключение имени события по типу вложения.")
    add_code_block(
        doc,
        """public Task SendReportPatchAsync(
    string groupKey, PatchReportSocketView view, string? signalRConnectionId)
    => SendToGroupAsync(groupKey, "ReceiveReportPatch", [view], signalRConnectionId);

public Task SendBugCreateAsync(
    string groupKey, BugSummaryDbModel summary, string? signalRConnectionId)
    => SendToGroupAsync(groupKey, "ReceiveBugCreate", [summary], signalRConnectionId);

public Task SendCommentCreateAsync(
    string groupKey, CommentSummaryDbModel comment, string? signalRConnectionId)
    => SendToGroupAsync(groupKey, "ReceiveCommentCreate", [comment], signalRConnectionId);

public Task SendAttachmentCreateAsync(
    string groupKey, AttachmentSocketView view, string? signalRConnectionId)
{
    string eventName = view.AttachType switch
    {
        (int)AttachType.Comment => "ReceiveCommentAttachmentCreate",
        (int)AttachType.BugStep => "ReceiveBugStepAttachmentCreate",
        _ => "ReceiveBugAttachmentCreate"
    };

    return SendToGroupAsync(groupKey, eventName, [view], signalRConnectionId);
}

public Task SendBugStepsOrderUpdateAsync(
    string groupKey, int bugId, BugStepSummaryDbModel[] steps,
    string? signalRConnectionId)
    => SendToGroupAsync(
        groupKey, "ReceiveBugStepsOrderUpdate", [bugId, steps], signalRConnectionId);""",
        "csharp",
        [],
        0,
    )

    p = doc.add_heading("ПРИЛОЖЕНИЕ В", level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.page_break_before = True
    add_centered(doc, "Фрагменты клиентского real-time слоя, проверочный клиент и результаты испытаний", bold=True)
    add_body_paragraph(doc, "Клиентская часть решения относится к real-time слою страницы отчета: построению SignalR WebSocket-соединения, типизированному контракту событий, регистрации единого набора обработчиков, обновлению диагностических данных, повторному вступлению в группу после восстановления соединения, доменной обработке событий и тестам этого поведения.")
    add_code_block(
        doc,
        """const reconnectDelays = [0, 2_000, 5_000, 10_000, 30_000, 60_000, 120_000];

export const buildConnection = (): HubConnection => {
  const wsPath = getAppWebSocketUrl("/report-page-hub", "v1");
  const fullUrl = `${window.location.origin}${wsPath}`;

  return new HubConnectionBuilder()
    .withUrl(fullUrl, {
      transport: HttpTransportType.WebSockets,
      skipNegotiation: import.meta.env.VITE_SIGNALR_SKIP_NEGOTIATION === "true",
    })
    .withAutomaticReconnect(reconnectDelays)
    .build();
};""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Контракт real-time событий типизирован: серверные имена событий перечислены в enum `SocketEvent`, а для каждого события задан тип полезной нагрузки в `SocketPayload`. Для событий, аргументы которых на стороне сервера передаются позиционно, заданы кастомные парсеры, преобразующие их в объект с именованными полями.")
    add_code_block(
        doc,
        """export enum SocketEvent {
  ReportPatch = "ReceiveReportPatch",
  ReportParticipant = "ReceiveReportParticipant",
  ReportLinkCreate = "ReceiveReportLinkCreate",
  ReportLinkUpdate = "ReceiveReportLinkUpdate",
  ReportLinkDelete = "ReceiveReportLinkDelete",
  BugCreate = "ReceiveBugCreate",
  BugPatch = "ReceiveBugPatch",
  CommentCreate = "ReceiveCommentCreate",
  CommentUpdate = "ReceiveCommentUpdate",
  CommentDelete = "ReceiveCommentDelete",
  BugStepCreate = "ReceiveBugStepCreate",
  BugStepPatch = "ReceiveBugStepPatch",
  BugStepsOrderUpdate = "ReceiveBugStepsOrderUpdate",
  BugStepDelete = "ReceiveBugStepDelete",
  BugAttachmentCreate = "ReceiveBugAttachmentCreate",
  // ... всего 24 события для страницы отчета
}

export type SocketPayload = {
  [SocketEvent.ReportPatch]: PatchReportSocketResponse;
  [SocketEvent.BugPatch]: { bugId: number; patch: PatchBugSocketResponse };
  [SocketEvent.CommentDelete]: { bugId: number; commentId: number };
  // ... соответствующие типы для остальных событий
};

export const customParsers: Partial<
  Record<SocketEvent, (...args: unknown[]) => SocketPayload[SocketEvent]>
> = {
  [SocketEvent.BugPatch]: (...args) => {
    const [bugId, patch] = args as [number, PatchBugSocketResponse];
    return { bugId, patch };
  },
  [SocketEvent.CommentDelete]: (...args) => {
    const [bugId, commentId] = args as [number, number];
    return { bugId, commentId };
  },
};""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Команды на серверный хаб оформлены как effector-эффекты. Это делает работу с SignalR частью реактивной модели приложения и позволяет наблюдать за состоянием подписки тем же способом, что и за остальными доменными операциями.")
    add_code_block(
        doc,
        """export const joinReportFx = socket.createEffect(
  async ({ conn, reportId }: { conn: HubConnection; reportId: string }) => {
    await conn.invoke("JoinReportGroupAsync", reportId);
  }
);

export const leaveReportFx = socket.createEffect(
  async ({ conn, reportId }: { conn: HubConnection; reportId: string }) => {
    await conn.invoke("LeaveReportGroupAsync", reportId);
  }
);""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Инициализация соединения построена как единая последовательность: создается `HubConnection`, по перечислению `SocketEvent` регистрируется единый набор обработчиков с поддержкой кастомных парсеров, подписки на системные события `onreconnected` и `onclose` обеспечивают восстановление и корректную очистку. Защита через `initPromise` исключает повторную параллельную инициализацию.")
    add_code_block(
        doc,
        """export const initSocketFx = socket.createEffect(async () => {
  const currentConn = $connection.getState();
  if (currentConn && currentConn.state !== HubConnectionState.Disconnected) {
    return;
  }

  if (initPromise) {
    await initPromise;
    return;
  }

  initPromise = (async () => {
    const conn = buildConnection();
    const handlers = new Map<SocketEvent, (p: unknown) => void>();

    Object.values(SocketEvent).forEach((event) => {
      const customParser = customParsers[event];
      const handler = (...args: unknown[]) => {
        const payload = customParser
          ? customParser(...args)
          : (args[0] as SocketPayload[SocketEvent]);
        socketEventReceived({ type: event, payload });
      };
      conn.on(event, handler);
      handlers.set(event, handler);
    });

    conn.onreconnected((connectionId) => {
      connectionReconnected();
      setSignalRConnectionId(connectionId ?? null);
      void refreshConnectionDiagnostics(conn);
    });

    conn.onclose((e) => {
      handlers.forEach((h, ev) => conn.off(ev, h));
      if ($connection.getState() === (conn as ConnectionReady)) {
        connectionClosed(e);
        setSignalRConnectionId(null);
      }
    });

    try {
      await startConnection(conn);
      connectionStarted(
        Object.assign(conn, { started: true }) as ConnectionReady
      );
      await refreshConnectionDiagnostics(conn);
    } catch (e) {
      handlers.forEach((h, ev) => conn.off(ev, h));
      connectionClosed(e as Error);
    }
  })();

  try { await initPromise; } finally { initPromise = null; }
});""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "После восстановления соединения клиент обновляет диагностические данные и сохраняет актуальный `connectionId`; это позволяет исключать собственное соединение при отправке события и проверять, к какому узлу подключен клиент после reconnect.")
    add_code_block(
        doc,
        """const refreshConnectionDiagnostics = async (conn: HubConnection) => {
  try {
    const diagnostics = await conn.invoke<ConnectionDiagnostics>(
      "GetConnectionDiagnosticsAsync"
    );
    setSignalRConnectionId(diagnostics.connectionId);
    connectionDiagnosticsUpdated(diagnostics);
  } catch (e) {
    console.error("[Socket] Failed to read connection diagnostics", e);
    connectionDiagnosticsUpdated(null);
  }
};""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Страница отчета хранит текущий `reportId` и после reconnect повторно вызывает `JoinReportGroupAsync`, чтобы восстановить membership в SignalR-группе. Логика повторной подписки выделена в отдельный модуль `createReconnectJoinHandler`, что упрощает тестирование.")
    add_code_block(
        doc,
        """const currentReportId = useRef<string | null>(null);

useEffect(() => {
  if (!connection) return;

  connection.onreconnected(
    createReconnectJoinHandler({
      join,
      getCurrentReportId: () => currentReportId.current,
    })
  );
}, [connection, join]);

export const createReconnectJoinHandler = ({ join, getCurrentReportId }) => {
  return () => {
    const reportId = getCurrentReportId();
    if (reportId != null) join(reportId);
  };
};""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Поведение повторного вступления в группу покрыто модульными тестами: проверяются отсутствие лишнего вызова без текущего отчета, повторная подписка после reconnect и использование актуального `reportId` при последовательных переподключениях.")
    add_code_block(
        doc,
        """it("does not call join when there is no current report id", () => {
  const calls: string[] = [];
  const handler = createReconnectJoinHandler({
    join: (reportId) => calls.push(reportId),
    getCurrentReportId: () => null,
  });

  handler();

  expect(calls).toEqual([]);
});

it("calls join with current report id after reconnect", () => {
  const calls: string[] = [];
  const handler = createReconnectJoinHandler({
    join: (reportId) => calls.push(reportId),
    getCurrentReportId: () => "42",
  });

  handler();

  expect(calls).toEqual(["42"]);
});

it("uses latest report id on each reconnect", () => {
  const calls: string[] = [];
  let currentReportId = "101";
  const handler = createReconnectJoinHandler({
    join: (reportId) => calls.push(reportId),
    getCurrentReportId: () => currentReportId,
  });

  handler();
  currentReportId = "202";
  handler();

  expect(calls).toEqual(["101", "202"]);
});""",
        "typescript",
        [],
        0,
    )
    add_body_paragraph(doc, "Полученные real-time события распределяются по доменным сторам страницы отчета: изменение отчета, создание и изменение бага, операции над комментариями, шагами воспроизведения и вложениями. Ниже приведена выжимка из 24 подписок хука `useReportSocketEvents`, демонстрирующая разложение real-time потока на доменные операции интерфейса.")
    add_code_block(
        doc,
        """export const useReportSocketEvents = () => {
  const reportId = useUnit($reportIdStore);
  const socketEvents = useUnit({
    patchReportSocketEvent,
    createBugSocketEvent,
    patchBugSocketEvent,
    createCommentSocketEvent,
    updateCommentSocketEvent,
    deleteCommentSocketEvent,
    bugAttachmentCreatedSocketEvent,
    // ...
  });

  useSocketEvent(SocketEvent.ReportPatch, (patch) =>
    socketEvents.patchReportSocketEvent(patch)
  );

  useSocketEvent(SocketEvent.BugCreate, (bug) => {
    if (!reportId) return;
    socketEvents.createBugSocketEvent({ reportId, bug });
  });

  useSocketEvent(SocketEvent.BugPatch, ({ bugId, patch }) =>
    socketEvents.patchBugSocketEvent({ bugId, patch })
  );

  useSocketEvent(SocketEvent.CommentCreate, (comment) =>
    socketEvents.createCommentSocketEvent(comment)
  );

  useSocketEvent(SocketEvent.CommentUpdate, (comment) =>
    socketEvents.updateCommentSocketEvent(comment)
  );

  useSocketEvent(SocketEvent.CommentDelete, ({ bugId, commentId }) =>
    socketEvents.deleteCommentSocketEvent({ bugId, commentId })
  );

  useSocketEvent(SocketEvent.BugAttachmentCreate, (attachment) => {
    if (attachment.attachType === AttachmentTypes.COMMENT) return;
    socketEvents.bugAttachmentCreatedSocketEvent(attachment);
  });
  // ... остальные подписки по аналогии
};""",
        "typescript",
        [],
        0,
    )
    doc.add_page_break()
    add_centered(doc, "Проверочный клиент", bold=True)
    add_body_paragraph(doc, "Чтобы изоляция от пользовательского интерфейса не мешала строгости проверки, проверочный клиент имеет собственные обвязки: единый таймаут ожидания событий, доменный HTTP-запрос на изменение отчета с пробросом `X-Signal-R-Connection-Id` (для исключения инициатора в `GroupExcept`), а также привязка соединения к конкретному серверному экземпляру по идентификатору `serverInstanceId`. Последнее особенно важно через nginx, который сам распределяет соединение между узлами.")
    add_code_block(
        doc,
        """const withTimeout = (promise, label) => {
  let timer;
  const timeout = new Promise((_, reject) => {
    timer = setTimeout(
      () => reject(new Error(`${label} timed out after ${timeoutMs}ms`)),
      timeoutMs,
    );
  });

  return Promise.race([promise, timeout]).finally(() => clearTimeout(timer));
};

const sendReportPatch = async (reportId, title, signalRConnectionId) =>
  requestJson(`${nodeA}/v2/reports/${reportId}`, {
    method: "PATCH",
    headers: { "X-Signal-R-Connection-Id": signalRConnectionId },
    body: JSON.stringify({ title }),
  });

const connectToServerInstance = async (baseUrl, serverInstanceId) => {
  for (let attempt = 1; attempt <= 8; attempt += 1) {
    const connection = buildConnection(baseUrl, { automaticReconnect: true });
    await connection.start();
    const diagnostics = await getDiagnostics(connection);

    if (diagnostics.serverInstanceId === serverInstanceId) {
      return { connection, diagnostics, attempt };
    }

    await connection.stop();
    await wait(150);
  }

  throw new Error(
    `Could not connect to ${serverInstanceId} through ${baseUrl}`,
  );
};""",
        "javascript",
        [],
        0,
    )
    add_body_paragraph(doc, "Базовый сценарий проверки доставки события подключает два SignalR-соединения к разным узлам, подписывает их на группу отчета и ожидает событие `ReceiveReportPatch` на втором узле. Замеры задержки выполняются по `performance.now()` между моментом отправки HTTP-запроса и моментом прихода события на втором клиенте.")
    add_code_block(
        doc,
        """const diagnosticsA = await getDiagnostics(clientA);
const diagnosticsB = await getDiagnostics(clientB);

await clientA.invoke("JoinReportGroupAsync", reportId);
await clientB.invoke("JoinReportGroupAsync", reportId);

const patchStartedAt = performance.now();
await sendReportPatch(reportId, patchedTitle, diagnosticsA.connectionId);
await withTimeout(patchWatcher.promise, "ReceiveReportPatch");

return {
  ok: patchWatcher.getReceivedPatch()?.title === patchedTitle,
  nodeA: diagnosticsA,
  nodeB: diagnosticsB,
  metrics: {
    deliveryLatencyMs: roundMetric(patchWatcher.getReceivedAt() - patchStartedAt),
  },
};""",
        "javascript",
        [],
        0,
    )
    add_body_paragraph(doc, "Сценарий проверки отказоустойчивости останавливает контейнер целевого узла и ожидает автоматический переход клиента на оставшийся узел через nginx. Подписывается обработчик `onreconnected`, который после восстановления соединения читает диагностику нового узла и повторно вызывает `JoinReportGroupAsync`. Сценарий считается пройденным, если событие изменения отчета доставлено клиенту на новом узле, а `serverInstanceId` после восстановления отличается от исходного.")
    add_code_block(
        doc,
        """const initial = await connectToServerInstance(nginxNode, failoverTargetInstance);
clientB = initial.connection;
const diagnosticsB = initial.diagnostics;

await clientB.invoke("JoinReportGroupAsync", reportId);

const reconnectStartedAt = performance.now();
const rejoinPromise = new Promise((resolve, reject) => {
  clientB.onreconnected(async () => {
    try {
      diagnosticsAfterFailover = await getDiagnostics(clientB);
      await clientB.invoke("JoinReportGroupAsync", reportId);
      resolve();
    } catch (error) {
      reject(error);
    }
  });
});

await runDocker("stop", "--time", "1", failoverContainer);
await withTimeout(rejoinPromise, "SignalR failover reconnect/rejoin");
const rejoinedAt = performance.now();

const secondPatchWatcher = waitForReportPatch(clientB, secondTitle);
await sendReportPatch(reportId, secondTitle, diagnosticsA.connectionId);
await withTimeout(secondPatchWatcher.promise, "Post-failover ReceiveReportPatch");

return {
  ok:
    secondPatchWatcher.getReceivedPatch()?.title === secondTitle &&
    diagnosticsAfterFailover?.serverInstanceId !== diagnosticsB.serverInstanceId,
  metrics: {
    failoverReconnectAndRejoinMs: roundMetric(rejoinedAt - reconnectStartedAt),
  },
};""",
        "javascript",
        [],
        0,
    )
    add_body_paragraph(doc, "Результаты испытаний фиксируют различие между режимом без backplane и режимом с Redis backplane.")
    add_code_block(
        doc,
        """Multi-instance без backplane:
  Error: ReceiveReportPatch timed out after 10000ms

Multi-instance с Redis backplane:
  ok: true
  nodeA.serverInstanceId: app-api-1
  nodeB.serverInstanceId: app-api-2
  deliveryLatencyMs: 9.5

THESIS_ITERATIONS=30:
  successful: 30
  failed: 0
  avg: 7.2 ms
  p50: 6.1 ms
  p95: 13.8 ms

THESIS_SCENARIO=rejoin:
  reconnectAndRejoinMs: 6.5
  deliveryLatencyMs: 15.2

THESIS_SCENARIO=failover:
  app-api-2 -> app-api-1
  failoverReconnectAndRejoinMs: 352.0
  deliveryLatencyMs: 5.4""",
        "text",
        [],
        0,
    )


def audit_docx(path: Path):
    with zipfile.ZipFile(path) as zf:
        names = set(zf.namelist())
        required = {
            "word/document.xml",
            "word/styles.xml",
            "word/numbering.xml",
            "[Content_Types].xml",
        }
        missing = required - names
        if missing:
            raise RuntimeError(f"Missing OOXML parts: {sorted(missing)}")
        text = zf.read("word/document.xml").decode("utf-8", errors="ignore")
        for marker in ["", "turn0", "TODO", "Deep Research", "ChatGPT"]:
            if marker in text:
                raise RuntimeError(f"Forbidden marker in document.xml: {marker}")


def build_document(kind: str = "vkr"):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    figure_paths = ensure_figures()

    if kind == "practice":
        doc = Document()
        setup_styles(doc)
        enable_field_update_on_open(doc)
        set_page_margins(doc.sections[0])
        doc.sections[0].footer.is_linked_to_previous = False

        add_toc(doc, kind=kind)

        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
        set_page_margins(body_section)
        set_start_page_number(body_section, 4)
        add_page_number_footer(body_section)
    else:
        doc = load_vkr_title_pages()
        setup_styles(doc)
        enable_field_update_on_open(doc)
        doc.sections[0].footer.is_linked_to_previous = False

        body_section = doc.add_section(WD_SECTION.NEW_PAGE)
        # add_section добавляет пустой параграф с sectPr в конец титульной секции.
        # Сжимаем его, иначе он выпадает на отдельную пустую страницу.
        sect_break_p = doc.paragraphs[-1]
        _freeze_paragraph_format(sect_break_p)
        sect_break_p.paragraph_format.line_spacing = Pt(1)
        set_page_margins(body_section)
        # ГОСТ-конвенция: титульный = 1, задание = 2 (без видимых номеров),
        # РЕФЕРАТ начинается с видимого "3".
        set_start_page_number(body_section, 3)
        # Шаблон оставил флаг <w:titlePg/> — он подавляет футер на первой странице
        # секции, из-за чего РЕФЕРАТ оставался без номера. Снимаем.
        body_sect_pr = body_section._sectPr
        for el in body_sect_pr.findall(qn("w:titlePg")):
            body_sect_pr.remove(el)
        add_page_number_footer(body_section)
        add_abstract(doc)
        add_toc(doc, kind=kind)

    add_abbreviations(doc, kind=kind)
    text = SOURCE.read_text(encoding="utf-8")
    if kind == "practice":
        text = transform_for_practice(text)
    process_markdown(doc, text, figure_paths, kind=kind)
    add_appendices(doc, kind=kind)

    out_docx = PRACTICE_OUT_DOCX if kind == "practice" else OUT_DOCX
    doc.save(out_docx)
    audit_docx(out_docx)
    print(out_docx)


def main(kind: str | None = None):
    selected = kind or ("practice" if "--practice" in sys.argv else "vkr")
    if selected not in {"vkr", "practice"}:
        raise ValueError(f"Unknown document kind: {selected}")
    build_document(selected)


if __name__ == "__main__":
    main()
